import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import re

load_dotenv()

# Chiede quali file analizzare
input_utente = input("Quali file CSV vuoi analizzare? (separali con una virgola) ")
lista_file = [f.strip() for f in input_utente.split(",")]

# Chiede cosa analizzare
tipo_analisi = input("Cosa vuoi analizzare? (es. profittabilità, trend, confronto segmenti) ")

# Crea la connessione con Claude
client = anthropic.Anthropic()

risultati = []

# Analizza ogni file
for nome_file in lista_file:
    print(f"\nAnalizzando {nome_file}...")
    
    with open(nome_file, "r") as file:
        csv_content = file.read()

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[
            {"role": "user", "content": f"""
            Analizza questo dataset di vendite con focus su: {tipo_analisi}
            
            Per ogni punto fornisci:
            - Osservazioni chiave
            - Numeri specifici dal dataset
            - Eventuali anomalie o pattern rilevanti
            
            Dati:
            {csv_content}
            """}
        ]
    )
    
    analisi = message.content[0].text
    risultati.append({"file": nome_file, "analisi": analisi})
    print(analisi)

# Chiede il nome del file di output
nome_output = input("\nCome vuoi chiamare il report? (senza estensione) ")

def is_table(lines, index):
    """Controlla se la riga corrente fa parte di una tabella Markdown"""
    return "|" in lines[index]

def aggiungi_tabella(doc, righe_tabella):
    """Converte righe Markdown in tabella Word reale"""
    # Filtra la riga separatore |---|---|
    righe_dati = [r for r in righe_tabella if not re.match(r'^\|[-|\s]+\|$', r.strip())]
    if not righe_dati:
        return
    
    # Estrae celle da ogni riga
    tabella_dati = []
    for riga in righe_dati:
        celle = [c.strip() for c in riga.strip().strip("|").split("|")]
        tabella_dati.append(celle)
    
    if not tabella_dati:
        return
    
    # Crea la tabella Word
    num_col = len(tabella_dati[0])
    table = doc.add_table(rows=len(tabella_dati), cols=num_col)
    table.style = "Table Grid"
    
    for i, riga in enumerate(tabella_dati):
        for j, cella in enumerate(riga):
            if j < num_col:
                cell = table.cell(i, j)
                cell.text = cella
                # Prima riga = intestazione in grassetto
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

def processa_testo(doc, testo):
    """Converte testo Markdown in formattazione Word"""
    righe = testo.split("\n")
    i = 0
    while i < len(righe):
        riga = righe[i]
        
        # Titoli
        if riga.startswith("### "):
            doc.add_heading(riga.replace("### ", "").strip(), level=3)
        elif riga.startswith("## "):
            doc.add_heading(riga.replace("## ", "").strip(), level=2)
        elif riga.startswith("# "):
            doc.add_heading(riga.replace("# ", "").strip(), level=1)
        
        # Tabelle
        elif "|" in riga:
            righe_tabella = []
            while i < len(righe) and "|" in righe[i]:
                righe_tabella.append(righe[i])
                i += 1
            aggiungi_tabella(doc, righe_tabella)
            continue
        
        # Righe vuote o separatori
        elif riga.strip() == "" or riga.strip() == "---":
            doc.add_paragraph("")
        
        # Testo normale con grassetto
        else:
            para = doc.add_paragraph()
            parti = re.split(r'\*\*(.*?)\*\*', riga)
            for idx, parte in enumerate(parti):
                run = para.add_run(parte)
                if idx % 2 == 1:
                    run.bold = True
        
        i += 1

# Salva tutto in un unico .docx
doc = Document()
doc.add_heading("Report Analisi Dataset", level=1)
doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
doc.add_paragraph(f"Focus analisi: {tipo_analisi}")

for r in risultati:
    doc.add_heading(f"Analisi: {r['file']}", level=2)
    processa_testo(doc, r["analisi"])

doc.save(f"{nome_output}.docx")
print(f"\nReport salvato in {nome_output}.docx!")